"""DOCX parser (python-docx).

Produces deterministic facts only — counts, headings, fonts, a text sample — never judgement.
Heading levels come from paragraph styles; equations from embedded OMML; figures from inline
shapes; tables from the table collection. Pages are estimated (DOCX has no fixed pagination)
from explicit/rendered page breaks, falling back to a word-count heuristic.
"""

from __future__ import annotations

import math
import re
from io import BytesIO

from docx import Document as load_docx
from docx.oxml.ns import qn

from app.parsers.models import Counts, Heading, ParsedDocument

_HEADING_RE = re.compile(r"^heading\s+([123])$", re.IGNORECASE)
_PROBLEM_RE = re.compile(r"^\s*(problem|exercise)\b", re.IGNORECASE)
_REFERENCE_HEADINGS = {"references", "bibliography", "works cited"}
_WORDS_PER_PAGE = 350
_SAMPLE_CHARS = 2000


class DocxParser:
    supported_types: tuple[str, ...] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    )

    def parse(self, content: bytes, *, filename: str | None = None) -> ParsedDocument:
        document = load_docx(BytesIO(content))

        headings: list[Heading] = []
        body_parts: list[str] = []
        references = 0
        problems = 0
        in_references = False

        for para in document.paragraphs:
            style_name = para.style.name if para.style is not None else ""
            level = _heading_level(style_name or "")
            text = para.text.strip()
            if level is not None:
                if text:
                    headings.append(Heading(level=level, text=text[:200]))
                in_references = text.lower() in _REFERENCE_HEADINGS
                continue
            if not text:
                continue
            body_parts.append(text)
            if in_references:
                references += 1
            if _PROBLEM_RE.match(text):
                problems += 1

        counts = Counts(
            pages=_estimate_pages(document, body_parts),
            figures=len(document.inline_shapes),
            tables=len(document.tables),
            equations=_count_equations(document),
            problems=problems,
            references=references,
        )
        sample = "\n".join(body_parts)[:_SAMPLE_CHARS]
        return ParsedDocument(
            counts=counts,
            headings=headings,
            language=_detect_language(document),
            fonts=_collect_fonts(document),
            raw_text_sample=sample or None,
        )


def _heading_level(style_name: str) -> str | None:
    match = _HEADING_RE.match(style_name.strip())
    return f"H{match.group(1)}" if match else None


def _count_equations(document) -> int:  # type: ignore[no-untyped-def]
    body = document.element.body
    return len(body.findall(".//" + qn("m:oMath")))


def _estimate_pages(document, body_parts: list[str]) -> int:  # type: ignore[no-untyped-def]
    body = document.element.body
    explicit = sum(
        1 for br in body.findall(".//" + qn("w:br")) if br.get(qn("w:type")) == "page"
    )
    rendered = len(body.findall(".//" + qn("w:lastRenderedPageBreak")))
    words = sum(len(part.split()) for part in body_parts)
    heuristic = math.ceil(words / _WORDS_PER_PAGE) if words else 0
    return max(explicit + rendered + 1, heuristic, 1)


def _detect_language(document) -> str | None:  # type: ignore[no-untyped-def]
    language = document.core_properties.language
    return language or None


def _collect_fonts(document) -> list[str]:  # type: ignore[no-untyped-def]
    fonts: set[str] = set()
    for para in document.paragraphs:
        for run in para.runs:
            if run.font is not None and run.font.name:
                fonts.add(run.font.name)
    return sorted(fonts)[:10]
