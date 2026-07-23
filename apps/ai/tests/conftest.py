"""Shared test fixtures — builds a real DOCX manuscript in memory."""

from __future__ import annotations

from collections.abc import Iterator
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from app.core.config import get_settings


@pytest.fixture(autouse=True)
def _hermetic_provider_settings(monkeypatch: pytest.MonkeyPatch) -> Iterator[None]:
    """Force AI_PROVIDER=mock for every test, regardless of what a developer's local .env
    contains.

    pydantic-settings reads apps/ai/.env (env_file=".env") whenever no real environment variable
    overrides a field. Once a developer sets a real AI_PROVIDER=gemini + GEMINI_API_KEY in their
    .env (exactly what Sprint 1's live smoke test instructs), a plain `pytest` run would silently
    pick that up too — every test asserting default/mock behavior would fail, and worse, tests
    that build the real AnalysisOrchestrator/AssistantOrchestrator would fire real, billed network
    calls against Gemini on every local test run. Setting a real env var here takes precedence over
    the .env file (matching python-dotenv/pydantic-settings precedence), so tests are hermetic
    regardless of the developer's local provider configuration.

    test_gemini_smoke.py is unaffected: it reads settings at import time (module collection, before
    any fixture runs), so its opt-in live check still sees the real .env key.
    """
    monkeypatch.setenv("AI_PROVIDER", "mock")
    get_settings.cache_clear()
    yield
    get_settings.cache_clear()


def _build_docx() -> bytes:
    document = Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("This chapter introduces the subject in detail. " * 20)
    document.add_heading("Methods", level=2)
    document.add_paragraph("Problem 1. Solve for x given the constraints.")
    document.add_paragraph("Exercise 2. Prove the stated theorem.")
    document.add_table(rows=2, cols=3)
    document.add_heading("References", level=1)
    document.add_paragraph("[1] Author A, A Foundational Text, 2019.")
    document.add_paragraph("[2] Author B, Another Reference, 2021.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_docx_bytes() -> bytes:
    return _build_docx()


@pytest.fixture
def sample_docx_url(tmp_path: Path) -> str:
    path = tmp_path / "manuscript.docx"
    path.write_bytes(_build_docx())
    return path.as_uri()  # file:///... URL, mirroring the local storage driver's signed URL


def _build_pdf() -> bytes:
    # A minimal, valid single-page PDF (US Letter). No fonts/images/tags — enough to exercise the
    # preflight pipeline end-to-end deterministically.
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    return _build_pdf()


@pytest.fixture
def sample_pdf_url(tmp_path: Path) -> str:
    path = tmp_path / "production.pdf"
    path.write_bytes(_build_pdf())
    return path.as_uri()


def _build_encrypted_pdf(*, user_password: str) -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt(user_password=user_password, owner_password="ownersecret123")
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


@pytest.fixture
def permission_only_pdf_bytes() -> bytes:
    """Encrypted with an empty user password — the common "restrict printing/editing" case that
    opens with no prompt in any real-world viewer. Must parse exactly like an unencrypted PDF."""
    return _build_encrypted_pdf(user_password="")


@pytest.fixture
def password_protected_pdf_bytes() -> bytes:
    """Encrypted with a real, non-empty user password — genuinely can't be opened without it."""
    return _build_encrypted_pdf(user_password="realsecret")
